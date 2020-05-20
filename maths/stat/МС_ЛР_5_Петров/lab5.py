from scipy.stats import norm, t, f, f_oneway, ttest_ind, ttest_rel
from math import sqrt, fabs
from scipy.stats import bartlett, levene
from statsmodels.formula.api import ols
import statsmodels.api as stm
import numpy as np
from sklearn.feature_selection import f_classif
import pandas as pd

def save_table_to_docx(data, file_name):
    import docx
    pass
    try:
        doc = docx.Document(file_name)
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    # table.style = 'TableGrid'
    i, j = 0, 0
    for row in data:
        for item in row:
            table.rows[i].cells[j % len(data[0])].text = str(item)
            j += 1
            if j % len(data[0]) == 0:
                j = 0
                i += 1
    doc.add_paragraph()
    doc.save(file_name)


def get_data_from_docx(file_name):
    from docx import Document
    doc = Document(file_name)
    data = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if cell.text != '':
                data.append(float(cell.text.replace(",", ".")))
    return data


def get_table_from_docx(file_name):
    from docx import Document
    doc = Document(file_name)
    data = [[],]
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if cell.text != '':
                data[-1].append(float(cell.text.replace(",", ".")))
        data.append([])
    return data


def first(data, a, sigma):
    alpha = 0.05
    norm1 = norm(0, 1)
    N = len(data)
    u1 = sum(data) / N
    f = norm1.pdf(1 - alpha)
    c = a
    c += sigma / (sqrt(N) * f)
    print("U1 = ", u1)
    print("C = ", c)
    table = []
    if u1 <= c:
        print("Принимаемая гипотеза - H0")
        table = [[' ', ' ', ' ', 'Принятая гипотеза'], ["%.6f" % u1, 0.05, "%.6f" % c, "H0"]]
    else:
        print("Принимаемая гипотеза - H1")
        table = [[' ', ' ', ' ', 'Принятая гипотеза'], ["%.6f" % u1, 0.05, "%.6f" % c, "H1"]]
    save_table_to_docx(table, "results1.docx")


def get_tnm(N, M, sr1, sr2, disp1, disp2):
    return ((sr1 - sr2) / sqrt(N * (disp1 - sr1**2) + M * (disp2 - sr2**2))) * sqrt(M * N * (N + M - 2) / (N + M))


def second(data1, data2, data3):
    # data1, data2, data3 = [], [], []
    # N = len(data) // 3
    # print(N)
    # for i in enumerate(data):
    #     if i[0] % 3 == 0:
    #         data1.append(i[1])
    #     elif i[0] % 3 == 1:
    #         data2.append(i[1])
    #     else:
    #         data3.append(i[1])

    sr1 = sum(data1) / len(data1)
    sr2 = sum(data2) / len(data2)
    sr3 = sum(data3) / len(data3)

    disp1 = sum([i**2 for i in data1]) / len(data1)
    disp2 = sum([i**2 for i in data2]) / len(data2)
    disp3 = sum([i**2 for i in data3]) / len(data3)

    s1 = (len(data1) / (len(data1) - 1)) * (disp1 - sr1**2)
    s2 = (len(data2) / (len(data2) - 1)) * (disp2 - sr2**2)
    s3 = (len(data3) / (len(data3) - 1)) * (disp3 - sr3**2)

    tnm12 = get_tnm(len(data1), len(data2), sr1, sr2, disp1, disp2)
    tnm13 = get_tnm(len(data1), len(data3), sr1, sr3, disp1, disp3)
    tnm23 = get_tnm(len(data2), len(data3), sr2, sr3, disp2, disp3)

    table = list()
    table.append(["%.6f" % sr1, "%.6f" % sr2, "%.6f" % disp1, "%.6f" % disp2, "%.6f" % s1, "%.6f" % s2, "%.6f" % tnm12])
    table.append(["%.6f" % sr1, "%.6f" % sr3, "%.6f" % disp1, "%.6f" % disp3, "%.6f" % s1, "%.6f" % s3, "%.6f" % tnm13])
    table.append(["%.6f" % sr2, "%.6f" % sr3, "%.6f" % disp2, "%.6f" % disp3, "%.6f" % s2, "%.6f" % s3, "%.6f" % tnm23])

    save_table_to_docx(table, 'results2.docx')

    table = [[], [], []]
    tinv = float("%.6f" % t.ppf(0.975, 2 * N - 2))
    table[0] = ["%.6f" % fabs(tnm12), tinv]
    table[0].append("ВЕРНА" if fabs(tnm12) <= tinv else "НЕВЕРНА")

    table[1] = ["%.6f" % fabs(tnm13), tinv]
    table[1].append("ВЕРНА" if fabs(tnm13) <= tinv else "НЕВЕРНА")

    table[2] = ["%.6f" % fabs(tnm23), tinv]
    table[2].append("ВЕРНА" if fabs(tnm23) <= tinv else "НЕВЕРНА")

    save_table_to_docx(table, 'results2.docx')


def third(data, data1, data2, data3):
    m = 3
    N = len(data) / m
    u = 0
    for i in data:
        u += i
    # print(u)
    u /= N*m
    # print(u)
    uj = [sum(data1) / len(data1), sum(data2) / len(data2), sum(data3) / len(data3)]
    so = 0
    sf = N * ((uj[0] - u)**2 + (uj[1] - u)**2 + (uj[2] - u)**2)
    for i in data:
        so += ((i - u)**2)
    s_ost = so - sf
    fnm = (sf / (m-1)) / (s_ost / (m*(N-1)))
    k1, k2 = m-1, m*(N-1)
    print("%.6f" % so)
    print("%.6f" % sf)
    print("%.6f" % s_ost)
    print("%.6f" % (sf / (m-1)))
    print("%.6f" % (s_ost / (m*(N-1))))
    print("%.6f" % k1)
    print("%.6f" % k2)
    print("%.6f" % fnm)

    print("z: ", f.ppf(0.95, k1, k2))  # hz
    print("Гипотеза принята" if fnm < f.ppf(0.95, k1, k2) else "Гипотеза НЕ принята")
    pass


def fourth(data, data1, data2, data3):
    print("Задача4: ")
    N = len(data1)
    m = 3
    z = f.ppf(0.975, N - 1, N - 1)
    print(z)

    x, y = sum(data1) / N, sum(data2) / N
    x2, y2 = sum([i**2 for i in data1]) / N, sum([i**2 for i in data2]) / N
    s2x, s2y = (N * (x2 - x**2)) / (N-1), (N * (y2 - y**2)) / (N-1)
    s1, s2 = max(s2x, s2y), min(s2x, s2y)
    k1 = k2 = N - 1
    fnm = s1 / s2
    print("%.6f" % s1, "%.6f" % s2, "%.6f" % k1, "%.6f" % k2, "%.6f" % fnm)
    print("Гипотеза принята" if fnm < z else "Гипотеза НЕ принята")

    x, y = sum(data1) / N, sum(data3) / N
    x2, y2 = sum([i ** 2 for i in data1]) / N, sum([i ** 2 for i in data3]) / N
    s2x, s2y = (N * (x2 - x ** 2)) / (N - 1), (N * (y2 - y ** 2)) / (N - 1)
    s1, s2 = max(s2x, s2y), min(s2x, s2y)
    k1 = k2 = N - 1
    fnm = s1 / s2
    print("%.6f" % s1, "%.6f" % s2, "%.6f" % k1, "%.6f" % k2, "%.6f" % fnm)
    print("Гипотеза принята" if fnm < z else "Гипотеза НЕ принята")

    x, y = sum(data2) / N, sum(data3) / N
    x2, y2 = sum([i ** 2 for i in data2]) / N, sum([i ** 2 for i in data3]) / N
    s2x, s2y = (N * (x2 - x ** 2)) / (N - 1), (N * (y2 - y ** 2)) / (N - 1)
    s1, s2 = max(s2x, s2y), min(s2x, s2y)
    k1 = k2 = N - 1
    fnm = s1 / s2
    print("%.6f" % s1, "%.6f" % s2, "%.6f" % k1, "%.6f" % k2, "%.6f" % fnm)
    print("Гипотеза принята" if fnm < z else "Гипотеза НЕ принята")
    pass


def fifth(data, data1, data2, data3):
    print(f_oneway(data1, data2, data3))
    print(ttest_ind(data1, data2))
    print(ttest_ind(data1, data3))
    print(ttest_ind(data2, data3))
    print(ttest_ind(data1, data2, equal_var=False))
    print(ttest_ind(data1, data3, equal_var=False))
    print(ttest_ind(data2, data3, equal_var=False))

    print(bartlett(data1, data2, data3))
    print(bartlett(data1, data2))
    print(bartlett(data1, data3))
    print(bartlett(data2, data3))

    z = f.ppf(0.975, N - 1, N - 1)
    print(z)


a = 6.1
sigma = 1.46
data = get_data_from_docx('1.docx')
table = get_table_from_docx('1.docx')
save_table_to_docx(table, 'results1.docx')
print(data)
first(data, a, sigma)
#
data = get_data_from_docx('2.docx')
data1, data2, data3 = [], [], []
N = len(data) // 3
print(N)
for i in enumerate(data):
    if i[0] % 3 == 0:
        data1.append(i[1])
    elif i[0] % 3 == 1:
        data2.append(i[1])
    else:
        data3.append(i[1])
print(data)

# table = get_table_from_docx('2.docx')
# save_table_to_docx(table, 'results2.docx')
# print(data)
# second(data1, data2, data3)

# third(data, data1, data2, data3)
# fourth(data, data1, data2, data3)
# print(data1)
# print(data2)
# print(data3)
# fifth(data, data1, data2, data3)
