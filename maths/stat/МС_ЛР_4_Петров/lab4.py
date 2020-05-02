from matplotlib import pyplot
from numpy import arange
from math import fabs, sqrt, log
from scipy.stats import kstest, ks_2samp
from scipy.stats import uniform


def get_emp_func(data, x):
    if x < data[0]:
        return 0
    elif x > data[-1]:
        return 1
    else:
        i = 0
        while x > data[i]:
            i += 1
        return i / len(data)


def first(data, a, b):
    N = len(data)
    fig = pyplot.figure()
    ax = fig.gca()
    ax.set_yticks(arange(0, 1.1, 0.1))
    pyplot.grid(True)
    x = list([i for i in arange(a, b, 0.01)])
    dn = -100
    prev_dn = -100
    for i in range(len(x) - 1):
        y = get_emp_func(data, x[i])
        pyplot.plot([x[i], x[i + 1]], [y, y], color='blue', linewidth=0.5)
    for i in range(len(x) - 1):
        y = (x[i] - a) / (b - a)
        pyplot.plot([x[i], x[i + 1]], [y, y], color='red', linewidth=0.5)
    pyplot.show()

    ans = [0 for i in range(9)]
    ans[0], ans[1], ans[2] = a, b, N
    for i in range(N):
        y = get_emp_func(data, data[i]) + 1 / N
        y2 = get_emp_func(data, data[i])
        Fx = ((data[i] - a) / (b - a))
        dn = max(dn, max(fabs(y - Fx), fabs(y2 - Fx)))
        if dn != prev_dn:
            ans[5] = data[i]
            ans[6] = Fx
            ans[7], ans[8] = y, y2
            prev_dn = dn
    ans[3], ans[4] = dn, dn * sqrt(N)
    print(ans)
    print(kstest(data, lambda param: uniform.cdf(param, loc=a, scale=b-a)))


def second(data1, data2, a, b):
    N = len(data1)
    M = len(data2)
    fig = pyplot.figure()
    ax = fig.gca()
    ax.set_yticks(arange(0, 1.1, 0.1))
    pyplot.grid(True)
    x = list([i for i in arange(a, b, 0.01)])
    for i in range(len(x) - 1):
        y = get_emp_func(data1, x[i])
        pyplot.plot([x[i], x[i + 1]], [y, y], color='red', linewidth=0.5)
    for i in range(len(x) - 1):
        y = get_emp_func(data2, x[i])
        pyplot.plot([x[i], x[i + 1]], [y, y], color='blue', linewidth=0.5)
    pyplot.show()

    dnm = -100
    prev_dnm = -100
    ans = [0 for i in range(9)]
    ans[0], ans[1] = N, M
    for j in range(N - 1):
        for k in range(M - 1):
            FN_xj, FN_xj_0 = (j + 1) / N, j / N
            FM_yk, FM_yk_0 = (k + 1) / M, k / M
            FM_xj = get_emp_func(data2, data1[j])
            FN_yk = get_emp_func(data1, data2[k])
            dnm = max([dnm, fabs(FN_xj - FM_xj), fabs(FN_xj_0 - FM_xj), fabs(FN_yk - FM_yk), fabs(FN_yk - FM_yk_0)])
            if dnm - prev_dnm > 0.01:
                ans[2] = dnm
                ans[3] = dnm * sqrt(N * M / (N + M))
                ans[4] = data1[j] if dnm == max(fabs(FN_xj - FM_xj), fabs(FN_xj_0 - FM_xj)) else data2[k]
                ans[5] = get_emp_func(data1, ans[4]) + 1 / N
                ans[6] = get_emp_func(data1, ans[4])
                ans[7] = get_emp_func(data2, ans[4]) + 1 / M
                ans[8] = get_emp_func(data2, ans[4]) + 1 / M
                prev_dnm = dnm
    print(ans)
    print(ks_2samp(data1, data2))


def third(data, a, b):
    N = len(data)
    ans = [0 for i in range(9)]
    ans[0], ans[1], ans[2] = a, b, N
    dn_plus = -100
    prev_dn_plus = -100
    Fx = lambda x: (x - a) / (b - a) if  a <= x < b else 0 if x < a else 1
    for j in range(1, N):
        fn_xj = j / N
        fn_xj_0 = (j - 1) / N
        dn_plus = max([dn_plus, fn_xj - Fx(data[j]), fn_xj_0 - Fx(data[j])])
        if dn_plus != prev_dn_plus:
            prev_dn_plus = dn_plus
            ans[3] = dn_plus
            ans[4] = ans[3]*sqrt(N)
            ans[5] = data[j]
            ans[6] = Fx(data[j])
            ans[7] = fn_xj
            ans[8] = fn_xj_0
    print(ans)
    s_alpha = sqrt(-0.5 * log(0.05))
    print(s_alpha)


def fourth(data, a, b):
    N = len(data)
    ans = [0 for i in range(9)]
    ans[0], ans[1], ans[2] = a, b, N
    dn_plus = -100
    prev_dn_plus = -100
    Fx = lambda x: (x - a) / (b - a) if a <= x < b else 0 if x < a else 1
    for j in range(1, N):
        fn_xj = j / N
        fn_xj_0 = (j - 1) / N
        dn_plus = max([dn_plus, Fx(data[j]) - fn_xj, Fx(data[j]) - fn_xj_0])
        if dn_plus != prev_dn_plus:
            prev_dn_plus = dn_plus
            ans[3] = dn_plus
            ans[4] = ans[3] * sqrt(N)
            ans[5] = data[j]
            ans[6] = Fx(data[j])
            ans[7] = fn_xj
            ans[8] = fn_xj_0
    print(ans)
    s_alpha = sqrt(-0.5 * log(0.05))
    print(s_alpha)


def get_data_from_docx(file_name):
    from docx import Document
    doc = Document(file_name)
    data = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if cell.text != '':
                data.append(float(cell.text.replace(",", ".")))
    return data


def save_table_to_docx(data, file_name):
    import docx
    pass
    try:
        doc = docx.Document(file_name)
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()
    table = doc.add_table(rows=len(data) // 10, cols=10)
    # table.style = 'TableGrid'
    i, j = 0, 0
    for item in data:
        table.rows[i].cells[j % 10].text = str(item)
        j += 1
        if j % 10 == 0:
            j = 0
            i += 1
    doc.add_paragraph()
    doc.save(file_name)

a = 5.4
b = 11.4
# a = 3.6
# b = 10.6

data1 = get_data_from_docx('UD-1.docx')
data2 = get_data_from_docx('UD-2.docx')
# save_table_to_docx(data1, "result1.docx")
# save_table_to_docx(sorted(data1), "result1.docx")
# save_table_to_docx(data2, "result2.docx")
# save_table_to_docx(sorted(data2), "result2.docx")

first(sorted(data1), a, b)
second(sorted(data1), sorted(data2), a, b)
third(sorted(data1), a, b)
fourth(sorted(data1), a, b)