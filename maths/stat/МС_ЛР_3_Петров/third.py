from math import log2, sqrt, fabs
from scipy.stats import norm
from matplotlib import pyplot
from numpy import arange
import docx


def get_table1(group, N, a, sigma):
    table = [["k", "ak", "ak - a / sigma", "1/phi", "Phi", "pk"]]
    arg = (group[0][0][0] - a) / sigma
    table.append([0, group[0][0][0], arg, norm.pdf(arg) / sigma, norm.cdf(arg), "-"])
    k = 1
    prev = 0
    for (i, j), nk, wk in group:
        pk = norm.cdf(j, a, sigma) - prev if k < len(group) - 1 else 1 - prev
        table.append([k, j, (j - a) / sigma, norm.pdf(j, a, sigma) / sigma, norm.cdf(j, a, sigma), pk])
        prev = norm.cdf(j, a, sigma)
        k += 1
    return table


def get_table2(group, N, a, sigma):
    table = [['k', "Интервал", "wk", "pk", "|wk - pk|", "N..."]]
    k = 1
    prev = 0
    for (i, j), ni, wi in group:
        pk = norm.cdf(j, a, sigma) - prev if k < 9 else 1 - prev
        table.append([k, "[{}, {}]".format(i, j), wi, pk, fabs(wi - pk), N * ((wi - pk)**2) / pk])
        prev = norm.cdf(j, a, sigma)
        k += 1
    return table


def get_a_sigma(group, N, h):
    a = 0
    sigma = 0
    for (i, j), nk, wk in group:
        xk = 0.5 * (i + j)
        a += xk * nk
        sigma += xk**2 * wk
    a /= N
    sigma = sigma - a ** 2 - (h ** 2) / 12
    return a, sigma


def draw(group, h, a, sigma):
    x = [(i + j) / 2 for (i, j), k, l in group]
    y = [wi / h for (_, _), _, wi in group]
    pyplot.bar(x, y, width=h, edgecolor='blue', color='white')
    x = [i for i in arange(group[0][0][0], group[-1][0][1], h / 50)]
    y = [norm.pdf((i - a) / sigma) / sigma for i in x]
    pyplot.plot(x, y, color='red')
    pyplot.show()


def save_table_to_docx(data, file_name):
    pass
    try:
        doc = docx.Document(file_name)
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    # table.style = 'TableGrid'
    i, j = 0, 0
    for row in data:
        for item in row:
            table.rows[i].cells[j].text = str(item)
            j += 1
        i += 1
        j = 0
    doc.add_paragraph()
    doc.save(file_name)


def first(data):
    group, h = get_group(data)
    print(group)
    N = sum([k for (i, j), k, l in group])
    a, sigma = get_a_sigma(group, N, h)
    print("a = ", a, "sigma = ", sigma)
    table1 = get_table1(group, N, a, sqrt(sigma))
    table2 = get_table2(group, N, a, sqrt(sigma))
    for i in table1:
        print(i)
    print()
    for i in table2:
        print(i)

    print(sum(i[4] for i in table2[1:]), sum([i[5] for i in table2[1:]]))
    tmp = [[]]
    for i in range(len(data)):
        tmp[-1].append(data[i])
        if i % 10 == 9:
            tmp.append([])
    save_table_to_docx(tmp, "3_1.docx")

    tmp = [[]]
    data.sort()
    for i in range(len(data)):
        tmp[-1].append(data[i])
        if i % 10 == 9:
            tmp.append([])
    save_table_to_docx(tmp, "3_1.docx")

    save_table_to_docx(group, "3_1.docx")

    save_table_to_docx(table1, "3_1.docx")
    save_table_to_docx(table2, "3_1.docx")

    # draw(group, h, a, sqrt(sigma))


def draw2(group, h, a, b):
    x = [(i + j) / 2 for (i, j), k, l in group]
    y = [wi / h for (_, _), _, wi in group]
    pyplot.bar(x, y, width=h, edgecolor='blue', color='white')
    x = [i for i in arange(a, b, h / 25)]
    y = [(1 / (b - a)) if a <= i <= b else 0 for i in x]
    pyplot.plot(x, y, color='red')
    pyplot.show()


def get_table3(group, N, m):
    table = [['k', "Интервал", "wk", "pk", "|wk - pk|", "N..."]]
    k = 1
    for (i, j), ni, wi in group:
        pk = 1 / m
        table.append([k, "[{}, {}]".format(i, j), wi, pk, fabs(wi - pk), N * ((wi - pk) ** 2) / pk])
        k += 1
    return table


def second(data):
    a = 6.4
    b = 12.4
    group, h = get_group(data, a, b)
    print(group)
    N = sum([k for (i, j), k, l in group])
    m = (1 + round(log2(N)))
    h = (b - a) / m
    table3 = get_table3(group, N, m)

    tmp = [[]]
    for i in range(len(data)):
        tmp[-1].append(data[i])
        if i % 10 == 9:
            tmp.append([])
    save_table_to_docx(tmp, "3_2.docx")

    tmp = [[]]
    data.sort()
    for i in range(len(data)):
        tmp[-1].append(data[i])
        if i % 10 == 9:
            tmp.append([])
    save_table_to_docx(tmp, "3_2.docx")

    save_table_to_docx(table3, "3_2.docx")
    save_table_to_docx(group, "3_2.docx")

    # draw2(group, h, a, b)



def get_group(data, a = None, b = None):
    N = len(data)
    m = 1 + round(log2(N))
    a0 = min(data) if a is None else a
    am = max(data) if b is None else b
    h = (am - a0) / m
    group = []
    for i in range(m):
        t = list(filter(lambda x: a0 + h*i <= x <= a0 + h*(i+1), data))
        group.append([(a0 + h*i, a0 + h*(i+1)), len(t), len(t)/N])
    return group, h


def get_data_from_docx(file_name):
    from docx import Document
    doc = Document(file_name)
    data = []
    for row in doc.tables[0].rows:
        for cell in row.cells:
            if cell.text != '':
                data.append(float(cell.text.replace(",", ".")))
    return data


data = get_data_from_docx('1.docx')
first(data)

data = get_data_from_docx('2.docx')
second(data)
