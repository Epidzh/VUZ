from numpy import arange
from matplotlib import pyplot
from math import sqrt, fabs, atan, exp
import docx


# функция f(x) для первого этапа
def f_x(x):
    return x*atan(x) + exp(-x)

def save_table_to_docx(data):
	file_name = "data.docx"
	try:
		doc = docx.Document(file_name)
	except docx.opc.exceptions.PackageNotFoundError:
		doc = docx.Document()
	table = doc.add_table(rows=len(data), cols=len(data[0]))
	table.style = 'Table Grid'
	i, j = 0, 0
	for row in data:
		for item in row:
			if item is float:
				table.rows[i].cells[j].text = "%.6f" % (item)
			else:
				table.rows[i].cells[j].text = str(item)
			j += 1
		i += 1
		j = 0
	doc.add_paragraph()
	doc.save(file_name)

# функция поиска минимума функции
def first():
    eps = 0.001
    a, b = -1, 0
    x = arange(a, b+0.01, 0.01)
    y = [f_x(i) for i in x]
    pyplot.grid(True)
    pyplot.plot(x, y, color="#0000AA", label='f(x)')

    F = (sqrt(5) - 1) / 2
    table = []
    while fabs(b - a) > eps:
        x1, x2 = a + (1 - F) * (b - a), a + F*(b - a)
        table.append([a, b, x1, x2])
        if f_x(x1) >= f_x(x2):
            a = x1
        else:
            b = x2
    print("Координаты минимума функции (точность=0.001): {}\n".format((a + b) / 2))
    print("Значение функции в точке минимума (точность=0.001): {}\n".format(f_x((a+b) / 2)))
    # for i in enumerate(table):
    # 	print("Шаг {} ; a={}, b={}, x1={}, x2={}".format(i[0], i[1][0], i[1][1], i[1][2], i[1][3]))
    # save_table_to_docx(table)
    return (a+b) / 2


# функция решения задачи Коши
def second(y1_a, y2_a, h):
	func_y1 = lambda x, y2: x**2 + 0.2*y2**2 # первая функция
	func_y2 = lambda x, y1, y2: x*y1 / y2 # вторая функция
	a, b = 0, 5 # границы отрезка

	x = arange(a, b+h, h) # генерируем сетку
	y1, y2 = [y1_a], [y2_a] # первые значения приближений

	table = []
	# генерация точек согласно формулам
	for point in enumerate(x[:-1]):
		y1.append(y1[-1] + h*func_y1(point[1], y2[-1]))
		y2.append(y2[-1] + h*func_y2(point[1], y1[-2], y2[-1]))
		table.append([point[0], point[1], round(y1[-1], 6), round(y2[-1], 6)])
	# save_table_to_docx(table)
	return [x, y1, y2]


if __name__ == "__main__": # ход решения
	# min_func_coord = first() # первый этап - поиск минимума
	pyplot.show()

	# второй этап - поиск решения задачи Коши (шаг=0.1)
	print('\n Шаг 0.1:')
	x, y1, y2 = second(min_func_coord, min_func_coord, 0.1)
	for i in enumerate(x):
		print('x={} , y1={}. y2={}'.format(round(i[1], 2), round(y1[i[0]], 5), round(y2[i[0]], 5)))
	pyplot.grid(True)
	pyplot.plot(x, y1, color="#AA0000", label='y1, h=0.1')
	pyplot.plot(x, y2, color="#00AA00", label='y2, h=0.1')

	# поиск решения задачи Коши (шаг=0.05)
	print('\n Шаг 0.05:')
	x, y1, y2 = second(min_func_coord, min_func_coord, 0.05)
	for i in enumerate(x):
		print('x={} , y1={}. y2={}'.format(round(i[1], 2), round(y1[i[0]], 5), round(y2[i[0]], 5)))
	pyplot.plot(x, y1, color="#0000AA", label='y1, h=0.05')
	pyplot.plot(x, y2, color="#FF00FF", label='y2, h=0.05')
	x = arange(-1, 5 + 0.01, 0.01)
	pyplot.plot(x, [f_x(i) for i in x], color="#FF00FF", label='f(x)')
	pyplot.legend()
	pyplot.show()