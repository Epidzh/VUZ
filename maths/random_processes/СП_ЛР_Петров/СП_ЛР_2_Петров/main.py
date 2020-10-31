import matplotlib.pyplot as plt
import numpy as np
import copy
import functools
from math import fabs
import docx

def save_table_to_docx(data):
	file_name = "data.docx"
	try:
		doc = docx.Document(file_name)
	except docx.opc.exceptions.PackageNotFoundError:
		doc = docx.Document()
	table = doc.add_table(rows=len(data), cols=len(data[0]))
	table.style = 'Table Grid'
	i, j = 0, 0
	for row in data:
		for item in row:
			if item is float:
				table.rows[i].cells[j].text = "%.6f" % (item)
			else:
				table.rows[i].cells[j].text = str(item)
			j += 1
		i += 1
		j = 0
	doc.add_paragraph()
	doc.save(file_name)


A = [[-3, 0, 1, 1, 1], [1, -3, 0, 1, 1], [0, 1, -2, 1, 0], [1, 1, 1, -3, 0], [0, 1, 1, 1, -3]]

tmp = copy.deepcopy(A)
tmp = np.transpose(copy.deepcopy(A))
tmp = np.append(tmp, [[1, 1, 1, 1, 1]], axis=0)
print(tmp)
r = np.linalg.solve(tmp[1:6], [0, 0, 0, 0, 1])
print(r)
print("sum: ", sum(r))
print(np.matmul(r, A))

ri_K = [0 for i in range(len(A))]
prev_ri_K = [0 for i in range(len(A))]
vi_K = [0 for i in range(len(A))]
ti_K = [0 for i in range(len(A))]

state_time = np.random.exponential(scale= -1 / A[0][0])
print(state_time)
t_sob = [state_time]

delta_k = None
K = 1
state = 0
table1 = []
while (delta_k is None) or (delta_k > 0.001):
	K += 1
	k_states = [i[0] for i in enumerate(A[state]) if i[1] > 0]
	lambda_i = -A[state][state]
	alpha = np.random.random_sample()
	# print("aplha"alpha)
	new_state = 0
	t = 1 / lambda_i
	while alpha > t:
		new_state += 1
		t += (1 / lambda_i)

	new_state = k_states[new_state]
	ri_K[new_state] = ri_K[new_state] + 1

	tau_l = np.random.exponential(scale= - A[new_state][new_state])
	t_sob.append(t_sob[-1] + tau_l)
	ti_K[new_state] += tau_l

	state = new_state
	# delta_k = max([fabs((ri_K[i] / K) - (prev_ri_K[i] / K)) for i in range(len(A))])
	# for i in range(len(ri_K)):
	# 	prev_ri_K[i] = ri_K[i]
	delta_k = max([fabs((ri_K[i] / K) - (ri_K[i] / (K - 1))) for i in range(len(A))])
	table1.append([K-1, round(t_sob[-2], 5), new_state + 1, round(tau_l, 5), round(delta_k, 5)])


delta_i_k = [ti_K[i] / ((t_sob[-1] - t_sob[0]) for i in range(len(A))]
print(delta_i_k)

vi_K = [i / (K-1) for i in ri_K]
print("rik: ", ri_K)
print("vik: ", vi_K)
print("tik: ", ti_K)
print("t sob: ", t_sob)
print(K)

table2 = []
for i in range(len(A)):
	table2.append([round(ri_K[i], 5), round(vi_K[i], 5), round(ti_K[i], 5), round(delta_i_k[i], 5)])

save_table_to_docx(table1)
save_table_to_docx(table2)