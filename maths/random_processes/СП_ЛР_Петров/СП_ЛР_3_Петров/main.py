import copy
import math
import numpy
import docx


def get_next_obj(table2, prev_event_time, deaths_nums):
	num, typee, death = None, None, None
	for obj_num, obj_type, birth_time, live_time, death_time, des1, des2 in table2:
		if not(obj_num in deaths_nums) and death_time > prev_event_time:
			if (death is None) or death_time < death:
				num, typee, death = obj_num, obj_type, death_time
	return death, num, typee


def save_table_to_docx(data):
	file_name = "data.docx"
	try:
		doc = docx.Document(file_name)
	except docx.opc.exceptions.PackageNotFoundError:
		doc = docx.Document()
	table = doc.add_table(rows=len(data), cols=len(data[0]))
	table.style = 'Table Grid'
	i, j = 0, 0
	for row in data:
		for item in row:
			if type(item) is float:
				table.rows[i].cells[j].text = str(round(item, 6))
			else:
				table.rows[i].cells[j].text = str(item)
			j += 1
		i += 1
		j = 0
	doc.add_paragraph()
	doc.save(file_name)


pn1, pn2, pm1, gamma1, gamma2 = 0.223, 0.521, 0.443, 0.21, 0.5
pn11 = 1 - pn1 - pn2
pm0 = 1 - pm1
table1, table2, table3, table4, table5 = [], [], [], [], []
events_num = 1


event_time = 0
system_state = [1, 0]
max_object_num = 1
event_type = "Sn(1)"
object_live_time = numpy.random.exponential(0.1)
wait_time = object_live_time
table1.append([events_num, event_time, event_type, object_live_time, -1, str(system_state), wait_time, 1, "N"])
table2.append([max_object_num, "N", 0, object_live_time, object_live_time, -1, -1])
deaths = []

while events_num < 100:
	events_num += 1
	event_time += table1[-1][-3]
	prev_event_type = table1[-1][2]
	death_object_num = table1[-1][-2]
	death_object_type = table1[-1][-1]
	deaths.append(death_object_num)


	if death_object_type == "N":
		w = numpy.random.random_sample()
		event_type = 0
		if w < pn1:
			event_type = "Sn(1)"
		elif w < pn1 + pn2:
			event_type = "Sn(2)"
		else:
			event_type = "Sn(3)"

		if event_type == "Sn(3)":
			max_object_num += 2
			tn = numpy.random.exponential(gamma1 * system_state[0] + 2 * gamma1 * system_state[1] + 0.1)
			tm = numpy.random.exponential(3 * gamma2 * system_state[0] + gamma2 * system_state[1])
			system_state[1] += 1
			table2[death_object_num-1][-2] = max_object_num - 1
			table2[death_object_num-1][-1] = max_object_num
			if tm < tn:
				table2.append([max_object_num - 1, "M", event_time, tm, event_time + tm, -1, -1])
				table2.append([max_object_num, "N", event_time, tn, event_time + tn, -1, -1])
				next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
				table1.append([events_num, event_time, "Sn(3)", tm, tn, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])
			else:
				table2.append([max_object_num - 1, "N", event_time, tn, event_time + tn, -1, -1])
				table2.append([max_object_num, "M", event_time, tm, event_time + tm, -1, -1])
				next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
				table1.append([events_num, event_time, "Sn(3)", tn, tm, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])



		elif event_type == "Sn(2)":
			max_object_num += 2
			tn1 = numpy.random.exponential(gamma1 * system_state[0] + 2 * gamma1 * system_state[1] + 0.1)
			tn2 = numpy.random.exponential(gamma1 * system_state[0] + 2 * gamma1 * system_state[1] + 0.1)
			system_state[0] += 1
			table2[death_object_num-1][-2] = max_object_num - 1
			table2[death_object_num-1][-1] = max_object_num
			if tn1 < tn2:
				table2.append([max_object_num - 1, "N", event_time, tn1, event_time + tn1, -1, -1])
				table2.append([max_object_num, "N", event_time, tn2, event_time + tn2, -1, -1])
				next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
				table1.append([events_num, event_time, "Sn(2)", tn1, tn2, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])
			else:
				table2.append([max_object_num-1, "N", event_time, tn2, event_time + tn2, -1, -1])
				table2.append([max_object_num, "N", event_time, tn1, event_time + tn1, -1, -1])
				next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
				table1.append([events_num, event_time, "Sn(2)", tn2, tn1, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])

		else:
			max_object_num += 1
			tn = numpy.random.exponential(gamma1 * system_state[0] + 2 * gamma1 * system_state[1] + 0.1)
			table2[death_object_num-1][-2] = max_object_num
			table2.append([max_object_num, "N", event_time, tn, event_time + tn, -1, -1])
			next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
			table1.append([events_num, event_time, "Sn(1)", tn, -1, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])
	else:
		w = numpy.random.random_sample()
		event_type = 0
		if w < pm1:
			event_type = "Sm(1)"
		else:
			event_type = "Sm(0)"

		if event_type == "Sm(1)":
			max_object_num += 1
			tm = numpy.random.exponential(3 * gamma2 * system_state[0] + gamma2 * system_state[1])
			table2[death_object_num-1][-2] = max_object_num
			table2.append([max_object_num, "M", event_time, tm, event_time + tm, -1, -1])
			next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
			table1.append([events_num, event_time, "Sm(1)", tm, -1, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])
		else:
			system_state[1] -= 1
			next_event_time, next_obj_num, next_obj_type = get_next_obj(table2, event_time, deaths)
			table1.append([events_num, event_time, "Sm(0)", -1, -1, str(system_state), next_event_time - event_time, next_obj_num, next_obj_type])

# print(table1)

for i in table1:
	print(i)
print("///////////////////////////////////////////////////////////////////////////////////////")

for i in table2:
	print(i)

t1 = [i[1] for i in table1[1:]]
t2 = list(sorted([i[4] for i in table2]))
count = 0
for i in range(len(t1)):
	if t1[i] == t2[i]:
		count += 1
	print(t1[i], '   ->   ', t2[i])
print(count)

table3.append([len(list(filter(lambda x: x[2] == "Sn(1)", table1)))])
table3[-1].append(len(list(filter(lambda x: x[2] == "Sn(2)", table1))))
table3[-1].append(len(list(filter(lambda x: x[2] == "Sn(3)", table1))))
table3[-1].append(len(list(filter(lambda x: x[2] == "Sm(0)", table1))))
table3[-1].append(len(list(filter(lambda x: x[2] == "Sm(1)", table1))))
table3[-1].append(sum(table3[0][:]))
table3.append([round(i / table3[0][-1], 6) for i in table3[0][:-1]])
table3[-1].append(sum(table3[1][:]))
print(table3)

table4.append([len(list(filter(lambda x: x[1] == "N", table2))), system_state[0]])
table4.append([len(list(filter(lambda x: x[1] == "M", table2))), system_state[1]])

print(table4)

states = {}
for i in table1:
	if i[5] in states.keys():
		states[i[5]] += 1
	else:
		states[i[5]] = 1
print(states)

sm1, sm2, sm3, sm4 = 0, 0, 0, 0
for i in states:
	time = 0
	for _, _, _, _, _, state, wait, _, _ in table1:
		if state == i:
			time += wait
	table5.append([i, states[i], states[i] / 100, time])
	table5[-1].append(table5[-1][-1] / table1[-1][1])
	sm1 += table5[-1][1]
	sm2 += table5[-1][2]
	sm3 += table5[-1][3]
	sm4 += table5[-1][4]
table5.append(["!", sm1, sm2, sm3, sm4])
for i in table5:
	print(i)
# print(table5)
save_table_to_docx(table1)
save_table_to_docx(table2)
save_table_to_docx(table3)
save_table_to_docx(table4)
save_table_to_docx(table5)