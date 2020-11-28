import numpy as np
from matplotlib import pyplot
import docx

def save_table_to_docx(data):
	file_name = "data.docx"
	try:
		doc = docx.Document(file_name)
	except docx.opc.exceptions.PackageNotFoundError:
		doc = docx.Document()
	table = doc.add_table(rows=len(data), cols=len(data[0]))
	table.style = 'Table Grid'
	i, j = 0, 0
	for row in data:
		for item in row:
			if type(item) is float:
				table.rows[i].cells[j].text = str(round(item, 6))
			else:
				table.rows[i].cells[j].text = str(item)
			j += 1
		i += 1
		j = 0
	doc.add_paragraph()
	doc.save(file_name)


def get_device_num(devices, t_sob):
	for device_num, device in enumerate(devices):
		if not (device[1] is None) and (device[1] - t_sob < 0.0000001):
			return device_num
	return -1

def get_free_device_num(devices):
	for device_num, device in enumerate(devices):
		if device[0] == False:
			return device_num
	return -1

control = 0
from math import fabs
while fabs(control - 1) > 0.00001:
	L, u = 7.18, 3.17

	count1, count2 = 0, 0
	K = 2
	table1, table2, table3 = [], [], []
	t_sob = 0
	state = 1
	devices = []

	t_sob = np.random.exponential(L)
	t_obsl = np.random.exponential(u)
	t_oz = np.random.exponential(L)
	t_min = t_obsl
	devices.append([True, t_sob + t_obsl])
	table1.append([K-1, t_sob, 1, state, t_min, t_oz, 1])
	table2.append([1, 1, t_obsl])

	while K < 1001:
		t1, t2 = table1[-1][-3], table1[-1][-2]
		if t1 == -1 or t1 > t2:  # новая заявка
			t_sob += t2
			t_obsl = np.random.exponential(u)
			state += 1
			device_num = get_free_device_num(devices)
			if device_num == -1:
				devices.append([True, t_sob + t_obsl])
				device_num = len(devices) - 1
				table2.append([len(devices), 1, t_obsl])
			else:
				devices[device_num] = [True, t_sob + t_obsl]
				table2[device_num][1] += 1
				table2[device_num][2] += t_obsl
			t_min = min([i - t_sob for st, i in devices if st is True])
			t_oz = np.random.exponential(L)
			table1.append([K, t_sob, 1, state, t_min, t_oz, device_num + 1])
			count1 += 1
		else:  # удаление заявки
			t_sob += t1
			state -= 1
			device_num = get_device_num(devices, t_sob)
			devices[device_num] = [False, None]
			t_min = -1
			if state > 0:
				t_min = min([i - t_sob for st, i in devices if st is True])
			table1.append([K, t_sob, 2, state, t_min, t2 - t1, device_num + 1])
			count2 += 1
		if table1[-1][4] < 0 and table1[-1][4] != -1:
			print(table1[-1])
		K += 1
	# print("/////////////////////////////////////////////////////////////////////////")
	# for i in table1:
	# 	print(i)
	# for i in table2:
	# 	print(i)
	devices_num = max([i[-1] for i in table1])
	# print(count1 + 1, count2)
	states_num = max([i[3] for i in table1])
	# print(states_num)
	for i in range(states_num + 1):
		v = len([j[3] for j in table1 if j[3] == i]) / 1000
		t = sum([j[-2] if j[-2] < j[-3] or j[-3] == -1 else j[-3] for j in table1 if j[3] == i])
		tau = t / table1[-1][1]
		table3.append([i, "r", v, tau])
	# for i in table3:
	# 	print(i)
	# print(sum([i[2] for i in table3]))
	control = sum([i[-1] for i in table3])
	print(control)
print(table1)
print(table2)
print(table3)
print(len(table1))
save_table_to_docx(table1[:100])
save_table_to_docx(table1[970:])
save_table_to_docx(table2)
save_table_to_docx(table3)