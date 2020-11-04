import numpy as np
from numpy.linalg import matrix_power
from numpy import matmul
from math import fabs
import copy
import docx


variant = 53
matrix = np.array([[0.272, 0.146, 0.582], [0, 0.701, 0.299], [0, 0.49, 0.51]])

def save_table_to_docx1(data, file_name):
    try:
        doc = docx.Document(file_name)
    except docx.opc.exceptions.PackageNotFoundError:
        doc = docx.Document()
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    i, j = 0, 0
    for row in data:
        for item in row:
            if item is float:
                table.rows[i].cells[j].text = item
            else:
                table.rows[i].cells[j].text = str(item)
            j += 1
        i += 1
        j = 0
    doc.add_paragraph()
    doc.save(file_name)


def convert_float(num):
    return "{:f}".format(float(num))


def getMatrixes1(matrix):
    ans = [matrix]
    deltas = ["-"]
    matrix_old = matrix
    delta = 0.00001
    delta_k = 0
    n = 1
    while delta_k > delta or n == 1:
        delta_k = 0
        n += 1
        matrix_new = matrix_power(matrix, n)
        ans.append(np.array(matrix_new))
        for i in range(len(matrix)):
            delta_k = max(delta_k, max([fabs(matrix_new[i][k] - matrix_old[i][k]) for k in range(len(matrix))]))
        matrix_old = matrix_new
        deltas.append("%.6f" % delta_k)

    for k in range(len(ans)):
        for i in range(len(ans[-1])):
            for j in range(len(ans[-1])):
                ans[k][i][j] = "%.6f" % (ans[k][i][j])
    # print(ans)
    return [ans, n, deltas, ans[-1]]


matrixes, n, deltas, final_matrix = getMatrixes1(matrix)
# print(n)

tmp = []
for index, item in enumerate(matrixes):
    tmp.append([index + 1])
    tmp[-1].append(item)
    tmp[-1].append(deltas[index])
# print(tmp)

save_table_to_docx1(tmp, "tables.docx")


# r = [46 / 71, 0, 25 / 71]
tmp = copy.deepcopy(matrix)
tmp = np.transpose(tmp)
for i in range(3):
    tmp[i][i] -= 1
tmp = np.append(tmp, [[1, 1, 1]], axis=0)
r = np.linalg.solve(tmp[1:4], [0,0,1])
print("r = ", r)


def getMatrixes3(matrix, vector, r):
    ans = [vector]
    n = 0
    delta = 0.00001
    delta_k = 0
    deltas = [max((1 - r[0]), r[1], r[2])]
    while delta_k > delta or n == 0:
        n += 1
        vector = matmul(vector, matrix)
        ans.append(list(vector))
        delta_k = max([fabs(vector[i] - r[i]) for i in range(3)])
        deltas.append(delta_k)
    return [ans, n, deltas]


vectors, m_min, deltas = getMatrixes3(matrix, [1, 0, 0], r)
print("100:")
print(vectors)
print(m_min)
print(deltas)

tmp = []
for index, item in enumerate(vectors):
    tmp.append([index])
    tmp[-1].append([round(i, 6) for i in item])
    tmp[-1].append(round(deltas[index], 6))
save_table_to_docx1(tmp, "tables.docx")

vectors, m_min, deltas = getMatrixes3(matrix, [0, 1, 0], r)
print("010:")
print(vectors)
print(m_min)
print(deltas)

tmp = []
for index, item in enumerate(vectors):
    tmp.append([index])
    tmp[-1].append([round(i, 6) for i in item])
    tmp[-1].append(round(deltas[index], 6))
save_table_to_docx1(tmp, "tables.docx")

vectors, m_min, deltas = getMatrixes3(matrix, [0, 0, 1], r)
print("001:")
print(vectors)
print(m_min)
print(deltas)

tmp = []
for index, item in enumerate(vectors):
    tmp.append([index])
    tmp[-1].append([round(i, 6) for i in item])
    tmp[-1].append(round(deltas[index], 6))
save_table_to_docx1(tmp, "tables.docx")


def task4(P, r):
    n_min = []
    for i in range(3):
        r_i_n = []
        v_i_n = []
        deltas = []
        p_0 = np.array([[1, 0, 0], [0, 1, 0], [0, 0, 1]])
        P_last = P
        n = 0
        R = 0
        delta = 1.0
        while delta > 0.001:
            n += 1
            tmp = p_0[i].dot(P_last)
            P_last = P_last.dot(P)
            alpha = np.random.random_sample()
            if alpha < tmp[0]:
                if i == 0:
                    R += 1
            elif alpha < tmp[0] + tmp[1]:
                if i == 1:
                    R += 1
            elif alpha <= 1:
                if i == 2:
                    R += 1
            v = R / n
            delta = fabs(v - r[i])
            r_i_n.append(R)
            v_i_n.append(round(v, 5))
            deltas.append(round(delta,5))

        print("V: ", v_i_n)
        print("R: ", r_i_n)
        n_min.append(n)
        tmp = []
        for index, item in enumerate(v_i_n):
            tmp.append([index+1])
            tmp[-1].append(r_i_n[index])
            tmp[-1].append(v_i_n[index])
            tmp[-1].append(deltas[index])

        save_table_to_docx1(tmp, "tables.docx")
    print('N_min: ', n_min)
task4(matrix, r)